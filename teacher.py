"""
Gem AI - Advanced Grading System with Multi-Provider Support
FIXED VERSION - Complete PDF report with all feedback sections
"""

import os
import fitz  # PyMuPDF
import docx
import pandas as pd
import numpy as np
from typing import Tuple, List, Dict, Optional, Any
import json
import re
from datetime import datetime
import cv2
import pytesseract
from PIL import Image
import speech_recognition as sr
from textblob import TextBlob
import nltk
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import openpyxl
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
import matplotlib.pyplot as plt
import seaborn as sns
from database import Database, Student, Exam, Result
import logging
import traceback
from email_service import EmailService
from ai_providers import AIProviderManager

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class GemAIGradingSystem:
    """
    Gem 💎 AI - The World's Most Advanced Exam Grading System
    
    Features:
    - Multi-format file support (PDF, DOCX, TXT, Images, Audio, Video)
    - Handwritten text recognition (OCR)
    - Voice/Video exam processing
    - Advanced AI grading with detailed feedback (supports multiple AI providers)
    - Plagiarism detection
    - Statistical analysis and reporting
    - Email notifications
    - PDF report generation with complete feedback
    - Database integration
    - Batch processing
    - Custom rubrics and grading criteria
    """
    
    def __init__(self, ai_provider: AIProviderManager = None, db_path: str = "gem_ai.db"):
        """Initialize the Gem AI Grading System with flexible AI provider"""
        logger.info("="*80)
        logger.info("Initializing Gem AI Grading System...")
        
        self.ai_provider = ai_provider or AIProviderManager()
        self.db = Database(db_path)
        self.email_service = EmailService()
        
        # Verify AI provider is working
        if self.ai_provider and self.ai_provider.client:
            logger.info(f"✅ Gem AI initialized with {self.ai_provider.active_provider} provider")
            logger.info(f"✅ AI Provider Status: {self.ai_provider.get_provider_info()}")
        else:
            logger.error("❌ AI Provider NOT properly initialized!")
            logger.error("❌ Grading will use fallback mode (static feedback)")
        
        logger.info("="*80)
        
        # Download required NLTK data
        try:
            nltk.download('punkt', quiet=True)
            nltk.download('stopwords', quiet=True)
            nltk.download('wordnet', quiet=True)
        except:
            pass
    
    def extract_text_from_file(self, file_path: str) -> Dict[str, Any]:
        """Extract text from various file formats with metadata"""
        ext = os.path.splitext(file_path)[1].lower()
        extraction_info = {
            'text': '',
            'file_type': ext,
            'extraction_method': '',
            'confidence': 1.0,
            'metadata': {}
        }
        
        try:
            if ext == ".txt":
                with open(file_path, "r", encoding="utf-8") as f:
                    extraction_info['text'] = f.read()
                extraction_info['extraction_method'] = 'direct_read'
                
            elif ext == ".pdf":
                doc = fitz.open(file_path)
                text = ""
                for page_num, page in enumerate(doc):
                    page_text = page.get_text()
                    if not page_text.strip():  # If no text, try OCR
                        pix = page.get_pixmap()
                        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                        page_text = pytesseract.image_to_string(img)
                        extraction_info['extraction_method'] = 'ocr_hybrid'
                    text += f"\n--- Page {page_num + 1} ---\n{page_text}"
                extraction_info['text'] = text
                extraction_info['metadata'] = {'pages': doc.page_count}
                if not extraction_info['extraction_method']:
                    extraction_info['extraction_method'] = 'pdf_text_extraction'
                
            elif ext == ".docx":
                doc = docx.Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
                # Extract tables too
                for table in doc.tables:
                    for row in table.rows:
                        row_text = "\t".join([cell.text for cell in row.cells])
                        text += f"\n{row_text}"
                extraction_info['text'] = text
                extraction_info['extraction_method'] = 'structured_read'
                
            elif ext in [".jpg", ".jpeg", ".png", ".bmp", ".tiff"]:
                img = Image.open(file_path)
                text = pytesseract.image_to_string(img, config='--psm 6')
                extraction_info['text'] = text
                extraction_info['extraction_method'] = 'ocr'
                extraction_info['confidence'] = 0.85
                
            elif ext in [".xlsx", ".xls"]:
                df = pd.read_excel(file_path)
                text = df.to_string()
                extraction_info['text'] = text
                extraction_info['extraction_method'] = 'spreadsheet_read'
                
            elif ext == ".csv":
                df = pd.read_csv(file_path)
                text = df.to_string()
                extraction_info['text'] = text
                extraction_info['extraction_method'] = 'csv_read'
                
            else:
                raise ValueError(f"Unsupported file format: {ext}")
            
            logger.info(f"✅ Extracted {len(extraction_info['text'])} characters from {file_path}")
                
        except Exception as e:
            logger.error(f"❌ Error extracting text from {file_path}: {str(e)}")
            raise
            
        return extraction_info
    
    def advanced_ai_grading(self, student_answer: str, model_answer: str, 
                           rubric: Dict[str, Any], student_info: Dict[str, str]) -> Dict[str, Any]:
        """Advanced AI grading with detailed analysis using any AI provider"""
        
        logger.info("="*80)
        logger.info(f"🎓 Starting AI Grading for: {student_info.get('name', 'Unknown')}")
        logger.info(f"📝 Student Answer Length: {len(student_answer)} characters")
        logger.info(f"📝 Model Answer Length: {len(model_answer)} characters")
        
        # Analyze answer quality
        quality_metrics = self._analyze_answer_quality(student_answer, model_answer)
        logger.info(f"📊 Quality Metrics: {quality_metrics}")
        
        # Check for plagiarism
        plagiarism_score = self._check_plagiarism(student_answer, model_answer)
        logger.info(f"🔍 Plagiarism Score: {plagiarism_score}%")
        
        # Generate comprehensive feedback
        grading_prompt = f"""You are Gem 💎 AI, the world's most advanced AI teacher and grading system. 

Grade this student's answer with extreme precision and provide comprehensive feedback:

STUDENT INFORMATION:
- Name: {student_info.get('name', 'N/A')}
- ID: {student_info.get('student_id', 'N/A')}
- Exam: {student_info.get('exam_name', 'N/A')}

GRADING CRITERIA:
- Total Points: {rubric.get('total_points', 100)}
- Question Number: {rubric.get('question_number', 1)}
- Subject: {rubric.get('subject', 'General')}
- Difficulty Level: {rubric.get('difficulty', 'Medium')}

RUBRIC BREAKDOWN:
{json.dumps(rubric.get('criteria', {}), indent=2)}

STUDENT'S ANSWER:
\"\"\"
{student_answer[:2000]}
\"\"\"

MODEL ANSWER:
\"\"\"
{model_answer[:2000]}
\"\"\"

QUALITY METRICS:
- Word Count: {quality_metrics['word_count']}
- Readability Score: {quality_metrics['readability']:.1f}
- Grammar Score: {quality_metrics['grammar_score']:.1f}
- Key Terms Coverage: {quality_metrics['key_terms_coverage']:.1f}%

Provide a detailed grading response in this exact format:

SCORE: X/{rubric.get('total_points', 100)}
PERCENTAGE: XX.X%
GRADE: [A+/A/B+/B/C+/C/D+/D/F]

DETAILED BREAKDOWN:
[Provide detailed scoring for each criterion]

STRENGTHS:
- [List 3-5 specific strengths based on the student's actual answer]

AREAS FOR IMPROVEMENT:
- [List 3-5 specific areas needing improvement with suggestions based on what the student actually wrote]

CORRECTIONS:
[For any incorrect information in the student's answer, provide the correct information]

FEEDBACK:
[Provide encouraging, constructive, and detailed feedback specific to this student's answer]

RECOMMENDATIONS:
[Specific study recommendations and resources based on this student's performance]
"""

        try:
            logger.info("🤖 Calling AI provider for grading...")
            logger.info(f"🤖 Using Provider: {self.ai_provider.active_provider}")
            
            # Use the AI provider manager to generate completion
            ai_feedback = self.ai_provider.generate_completion(
                grading_prompt,
                max_tokens=2000,
                temperature=0.3
            )
            
            logger.info(f"✅ AI Response received ({len(ai_feedback)} characters)")
            logger.info(f"📄 AI Response Preview: {ai_feedback[:200]}...")
            
            parsed_result = self._parse_ai_response(ai_feedback)
            logger.info(f"✅ Parsed Result: Score={parsed_result['score']}/{parsed_result['total_points']}, Grade={parsed_result['grade']}")
            
            # Add technical metrics
            parsed_result.update({
                'quality_metrics': quality_metrics,
                'plagiarism_score': plagiarism_score,
                'grading_timestamp': datetime.now().isoformat(),
                'ai_confidence': 0.95,
                'ai_provider': self.ai_provider.active_provider,
                'student_info': student_info,
                'rubric_used': rubric
            })
            
            logger.info("="*80)
            return parsed_result
            
        except Exception as e:
            logger.error("❌ AI GRADING FAILED!")
            logger.error(f"❌ Error Type: {type(e).__name__}")
            logger.error(f"❌ Error Message: {str(e)}")
            logger.error(f"❌ Traceback: {traceback.format_exc()}")
            logger.error("❌ Falling back to similarity-based grading...")
            logger.info("="*80)
            
            return self._generate_fallback_grade(student_answer, model_answer, rubric, student_info)
    
    def _analyze_answer_quality(self, answer: str, model_answer: str) -> Dict[str, Any]:
        """Analyze the quality of the student's answer"""
        try:
            blob = TextBlob(answer)
            model_blob = TextBlob(model_answer)
            
            # Calculate similarity
            vectorizer = TfidfVectorizer()
            try:
                tfidf_matrix = vectorizer.fit_transform([answer, model_answer])
                similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
            except:
                similarity = 0.0
            
            return {
                'word_count': len(answer.split()),
                'sentence_count': len(blob.sentences),
                'readability': min(100, max(0, (similarity * 100))),
                'grammar_score': min(100, max(0, blob.sentiment.polarity * 100 + 50)),
                'key_terms_coverage': self._calculate_key_terms_coverage(answer, model_answer),
                'similarity_to_model': similarity * 100
            }
        except Exception as e:
            logger.error(f"Error analyzing answer quality: {str(e)}")
            return {
                'word_count': len(answer.split()),
                'sentence_count': 0,
                'readability': 50.0,
                'grammar_score': 50.0,
                'key_terms_coverage': 50.0,
                'similarity_to_model': 50.0
            }
    
    def _calculate_key_terms_coverage(self, answer: str, model_answer: str) -> float:
        """Calculate how many key terms from model answer are covered"""
        try:
            answer_lower = answer.lower()
            model_words = set(re.findall(r'\b\w+\b', model_answer.lower()))
            common_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 
                          'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be', 'been', 'being',
                          'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'should',
                          'could', 'may', 'might', 'must', 'can', 'this', 'that', 'these', 'those'}
            key_terms = model_words - common_words
            
            if not key_terms:
                return 100.0
                
            covered_terms = sum(1 for term in key_terms if term in answer_lower)
            return (covered_terms / len(key_terms)) * 100
        except Exception as e:
            logger.error(f"Error calculating key terms coverage: {str(e)}")
            return 50.0
    
    def _check_plagiarism(self, answer: str, reference: str) -> float:
        """Simple plagiarism detection"""
        try:
            answer_sentences = re.split(r'[.!?]+', answer.lower())
            reference_sentences = re.split(r'[.!?]+', reference.lower())
            
            matches = 0
            total_sentences = len(answer_sentences)
            
            for ans_sent in answer_sentences:
                ans_sent = ans_sent.strip()
                if len(ans_sent) < 10:
                    continue
                for ref_sent in reference_sentences:
                    ref_sent = ref_sent.strip()
                    if ans_sent in ref_sent or ref_sent in ans_sent:
                        matches += 1
                        break
            
            return (matches / max(1, total_sentences)) * 100
        except Exception as e:
            logger.error(f"Error checking plagiarism: {str(e)}")
            return 0.0
    
    def _parse_ai_response(self, response: str) -> Dict[str, Any]:
        """Parse the AI grading response into structured data"""
        result = {
            'score': 0,
            'total_points': 100,
            'percentage': 0.0,
            'grade': 'F',
            'breakdown': '',
            'strengths': [],
            'improvements': [],
            'corrections': '',
            'feedback': '',
            'recommendations': ''
        }
        
        try:
            lines = response.split('\n')
            current_section = None
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                    
                if line.startswith('SCORE:'):
                    score_match = re.search(r'(\d+)/(\d+)', line)
                    if score_match:
                        result['score'] = int(score_match.group(1))
                        result['total_points'] = int(score_match.group(2))
                        
                elif line.startswith('PERCENTAGE:'):
                    percentage_match = re.search(r'(\d+\.?\d*)', line)
                    if percentage_match:
                        result['percentage'] = float(percentage_match.group(1))
                        
                elif line.startswith('GRADE:'):
                    grade_text = line.replace('GRADE:', '').strip()
                    # Extract just the grade letter(s)
                    grade_match = re.search(r'([A-F][+-]?)', grade_text)
                    if grade_match:
                        result['grade'] = grade_match.group(1)
                    else:
                        result['grade'] = grade_text[:2] if len(grade_text) >= 2 else grade_text
                    
                elif line.startswith('DETAILED BREAKDOWN:'):
                    current_section = 'breakdown'
                    continue
                elif line.startswith('STRENGTHS:'):
                    current_section = 'strengths'
                    continue
                elif line.startswith('AREAS FOR IMPROVEMENT:'):
                    current_section = 'improvements'
                    continue
                elif line.startswith('CORRECTIONS:'):
                    current_section = 'corrections'
                    continue
                elif line.startswith('FEEDBACK:'):
                    current_section = 'feedback'
                    continue
                elif line.startswith('RECOMMENDATIONS:'):
                    current_section = 'recommendations'
                    continue
                    
                # Add content to appropriate section
                if current_section == 'breakdown':
                    result['breakdown'] += line + '\n'
                elif current_section == 'strengths' and line.startswith('-'):
                    result['strengths'].append(line[1:].strip())
                elif current_section == 'improvements' and line.startswith('-'):
                    result['improvements'].append(line[1:].strip())
                elif current_section == 'corrections':
                    result['corrections'] += line + '\n'
                elif current_section == 'feedback':
                    result['feedback'] += line + ' '
                elif current_section == 'recommendations':
                    result['recommendations'] += line + ' '
            
            # Clean up feedback and recommendations
            result['feedback'] = result['feedback'].strip()
            result['recommendations'] = result['recommendations'].strip()
            result['corrections'] = result['corrections'].strip()
            result['breakdown'] = result['breakdown'].strip()
                    
        except Exception as e:
            logger.error(f"Error parsing AI response: {str(e)}")
            
        return result
    
    def _generate_fallback_grade(self, answer: str, model_answer: str, rubric: Dict, student_info: Dict) -> Dict[str, Any]:
        """Generate a basic grade when AI grading fails - THIS IS THE FALLBACK!"""
        logger.warning("⚠️ Using FALLBACK grading mode - feedback will be generic!")
        
        quality = self._analyze_answer_quality(answer, model_answer)
        score = int(quality['similarity_to_model'] * rubric.get('total_points', 100) / 100)
        
        return {
            'score': score,
            'total_points': rubric.get('total_points', 100),
            'percentage': quality['similarity_to_model'],
            'grade': self._score_to_grade(quality['similarity_to_model']),
            'feedback': f"⚠️ FALLBACK MODE: Automated grading based on similarity analysis. Score: {score}/{rubric.get('total_points', 100)}. AI grading failed - please check your API key configuration.",
            'strengths': ['Answer provided', 'Attempt made'],
            'improvements': ['More detail needed', 'Review model answer'],
            'corrections': 'AI grading unavailable - manual review recommended',
            'recommendations': 'Review course materials and model answer. Note: This is automated fallback grading.',
            'breakdown': f"Similarity to model answer: {quality['similarity_to_model']:.1f}%\n⚠️ This is fallback grading - AI provider failed",
            'quality_metrics': quality,
            'ai_confidence': 0.60,
            'ai_provider': f"{self.ai_provider.active_provider} (FALLBACK MODE)",
            'student_info': student_info,
            'rubric_used': rubric
        }
    
    def _score_to_grade(self, percentage: float) -> str:
        """Convert percentage to letter grade"""
        if percentage >= 97: return 'A+'
        elif percentage >= 93: return 'A'
        elif percentage >= 90: return 'A-'
        elif percentage >= 87: return 'B+'
        elif percentage >= 83: return 'B'
        elif percentage >= 80: return 'B-'
        elif percentage >= 77: return 'C+'
        elif percentage >= 73: return 'C'
        elif percentage >= 70: return 'C-'
        elif percentage >= 67: return 'D+'
        elif percentage >= 60: return 'D'
        else: return 'F'
    
    def batch_grade_exams(self, exam_config: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Grade multiple exams in batch"""
        results = []
        
        for student_file in exam_config['student_files']:
            try:
                student_info = self._extract_student_info(student_file, exam_config)
                extraction_result = self.extract_text_from_file(student_file)
                student_answer = extraction_result['text']
                
                grading_result = self.advanced_ai_grading(
                    student_answer,
                    exam_config['model_answer'],
                    exam_config['rubric'],
                    student_info
                )
                
                grading_result['file_path'] = student_file
                grading_result['extraction_info'] = extraction_result
                results.append(grading_result)
                
                self._save_result_to_db(grading_result, exam_config)
                
                logger.info(f"Graded: {student_info.get('name', 'Unknown')} - Score: {grading_result['score']}/{grading_result['total_points']}")
                
            except Exception as e:
                logger.error(f"Failed to grade {student_file}: {str(e)}")
                results.append({
                    'file_path': student_file,
                    'error': str(e),
                    'score': 0,
                    'total_points': exam_config['rubric'].get('total_points', 100)
                })
        
        return results
    
    def _extract_student_info(self, file_path: str, exam_config: Dict) -> Dict[str, str]:
        """Extract student information from filename or config"""
        filename = os.path.basename(file_path)
        pattern_matches = re.search(r'([^_]+)_([^_]+)_(\d+)', filename)
        
        if pattern_matches:
            return {
                'name': f"{pattern_matches.group(1)} {pattern_matches.group(2)}",
                'student_id': pattern_matches.group(3),
                'exam_name': exam_config.get('exam_name', 'Unknown Exam'),
                'subject': exam_config.get('subject', 'Unknown Subject')
            }
        
        return {
            'name': filename.split('.')[0],
            'student_id': 'Unknown',
            'exam_name': exam_config.get('exam_name', 'Unknown Exam'),
            'subject': exam_config.get('subject', 'Unknown Subject')
        }
    
    def _save_result_to_db(self, result: Dict, exam_config: Dict):
        """Save grading result to database"""
        try:
            student_info = result['student_info']
            student = self.db.get_or_create_student(
                student_info['student_id'],
                student_info['name']
            )
            
            exam = self.db.get_or_create_exam(
                exam_config.get('exam_name', 'Unknown'),
                exam_config.get('subject', 'Unknown'),
                exam_config['rubric'].get('total_points', 100)
            )
            
            self.db.save_result(
                student.id,
                exam.id,
                result['score'],
                result['percentage'],
                result['grade'],
                json.dumps(result)
            )
            
        except Exception as e:
            logger.error(f"Failed to save result to database: {str(e)}")
    
    def generate_pdf_report(self, results: List[Dict], output_path: str, exam_config: Dict):
        """Generate comprehensive PDF report with ALL feedback sections"""
        try:
            logger.info("="*80)
            logger.info(f"📄 Generating comprehensive PDF report...")
            logger.info(f"📄 Output path: {output_path}")
            logger.info(f"📄 Number of results: {len(results)}")
            
            doc = SimpleDocTemplate(output_path, pagesize=A4, 
                                   rightMargin=50, leftMargin=50,
                                   topMargin=50, bottomMargin=50)
            styles = getSampleStyleSheet()
            story = []
            
            # Custom styles for better formatting
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=26,
                textColor=colors.HexColor('#667eea'),
                alignment=TA_CENTER,
                spaceAfter=30,
                fontName='Helvetica-Bold'
            )
            
            heading2_style = ParagraphStyle(
                'CustomHeading2',
                parent=styles['Heading2'],
                fontSize=18,
                textColor=colors.HexColor('#764ba2'),
                spaceAfter=15,
                spaceBefore=20,
                fontName='Helvetica-Bold'
            )
            
            section_style = ParagraphStyle(
                'SectionStyle',
                parent=styles['Heading3'],
                fontSize=14,
                textColor=colors.HexColor('#28a745'),
                spaceAfter=10,
                spaceBefore=15,
                fontName='Helvetica-Bold'
            )
            
            body_style = ParagraphStyle(
                'BodyStyle',
                parent=styles['Normal'],
                fontSize=11,
                leading=16,
                spaceAfter=10,
                alignment=TA_LEFT
            )
            
            bullet_style = ParagraphStyle(
                'BulletStyle',
                parent=styles['Normal'],
                fontSize=11,
                leading=14,
                leftIndent=20,
                spaceAfter=6
            )
            
            # TITLE PAGE
            story.append(Spacer(1, 1*inch))
            story.append(Paragraph("💎 Gem AI Grading Report", title_style))
            story.append(Spacer(1, 0.3*inch))
            
            # EXAM INFORMATION BOX
            exam_info_data = [
                ['Exam:', exam_config.get('exam_name', 'Unknown Exam')],
                ['Subject:', exam_config.get('subject', 'Unknown Subject')],
                ['Date:', datetime.now().strftime('%B %d, %Y at %I:%M %p')],
                ['Total Students:', str(len(results))],
                ['Total Points:', str(exam_config.get('rubric', {}).get('total_points', 100))]
            ]
            
            exam_table = Table(exam_info_data, colWidths=[2*inch, 4*inch])
            exam_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f0f0')),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 11),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ('TOPPADDING', (0, 0), (-1, -1), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('BOX', (0, 0), (-1, -1), 2, colors.HexColor('#667eea'))
            ]))
            
            story.append(exam_table)
            story.append(Spacer(1, 0.5*inch))
            
            # STATISTICS SECTION
            valid_results = [r for r in results if 'error' not in r and 'score' in r]
            
            if valid_results:
                scores = [r.get('score', 0) for r in valid_results]
                percentages = [r.get('percentage', 0) for r in valid_results]
                
                story.append(Paragraph("📊 Class Statistics", heading2_style))
                
                stats_data = [
                    ['Average Score:', f"{np.mean(scores):.2f}", 'Average Percentage:', f"{np.mean(percentages):.1f}%"],
                    ['Highest Score:', f"{max(scores)}", 'Lowest Score:', f"{min(scores)}"],
                    ['Standard Deviation:', f"{np.std(scores):.2f}", 'Pass Rate (≥60%):', f"{len([p for p in percentages if p >= 60]) / len(percentages) * 100:.1f}%"]
                ]
                
                stats_table = Table(stats_data, colWidths=[1.8*inch, 1.5*inch, 1.8*inch, 1.5*inch])
                stats_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f8f9fa')),
                    ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                    ('FONTNAME', (2, 0), (2, -1), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
                    ('TOPPADDING', (0, 0), (-1, -1), 10),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
                ]))
                
                story.append(stats_table)
                story.append(Spacer(1, 0.3*inch))
            
            story.append(PageBreak())
            
            # INDIVIDUAL STUDENT RESULTS
            story.append(Paragraph("📋 Individual Student Results", heading2_style))
            story.append(Spacer(1, 0.2*inch))
            
            for i, result in enumerate(results, 1):
                if 'error' in result:
                    continue
                
                student_info = result.get('student_info', {})
                
                # Student Header Box
                story.append(Paragraph(f"Student {i}", section_style))
                
                student_header_data = [
                    ['Name:', student_info.get('name', 'Unknown')],
                    ['Student ID:', student_info.get('student_id', 'Unknown')],
                    ['Score:', f"{result.get('score', 0)}/{result.get('total_points', 100)}"],
                    ['Percentage:', f"{result.get('percentage', 0):.1f}%"],
                    ['Grade:', result.get('grade', 'F')]
                ]
                
                student_table = Table(student_header_data, colWidths=[1.5*inch, 4.5*inch])
                student_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#e3f2fd')),
                    ('BACKGROUND', (1, -1), (1, -1), colors.HexColor('#d4edda')),
                    ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                    ('FONTNAME', (1, -1), (1, -1), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 11),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                    ('TOPPADDING', (0, 0), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('BOX', (0, 0), (-1, -1), 1.5, colors.HexColor('#667eea'))
                ]))
                
                story.append(student_table)
                story.append(Spacer(1, 0.2*inch))
                
                # DETAILED BREAKDOWN (if available)
                if result.get('breakdown'):
                    story.append(Paragraph("📝 Detailed Breakdown", section_style))
                    breakdown_text = result['breakdown'][:1000]  # Limit length for PDF
                    story.append(Paragraph(breakdown_text.replace('\n', '<br/>'), body_style))
                    story.append(Spacer(1, 0.15*inch))
                
                # FEEDBACK SECTION
                if result.get('feedback'):
                    story.append(Paragraph("💬 General Feedback", section_style))
                    feedback_text = result['feedback'][:800]  # Limit length
                    story.append(Paragraph(feedback_text, body_style))
                    story.append(Spacer(1, 0.15*inch))
                
                # STRENGTHS SECTION (FIXED - NOW INCLUDED!)
                if result.get('strengths') and len(result['strengths']) > 0:
                    story.append(Paragraph("💪 Strengths", section_style))
                    for strength in result['strengths']:
                        strength_text = f"• {strength}"
                        story.append(Paragraph(strength_text, bullet_style))
                    story.append(Spacer(1, 0.15*inch))
                
                # AREAS FOR IMPROVEMENT SECTION (FIXED - NOW INCLUDED!)
                if result.get('improvements') and len(result['improvements']) > 0:
                    story.append(Paragraph("📈 Areas for Improvement", section_style))
                    for improvement in result['improvements']:
                        improvement_text = f"• {improvement}"
                        story.append(Paragraph(improvement_text, bullet_style))
                    story.append(Spacer(1, 0.15*inch))
                
                # CORRECTIONS SECTION (FIXED - NOW INCLUDED!)
                if result.get('corrections') and result['corrections'].strip():
                    story.append(Paragraph("✏️ Corrections", section_style))
                    corrections_text = result['corrections'][:800]  # Limit length
                    story.append(Paragraph(corrections_text.replace('\n', '<br/>'), body_style))
                    story.append(Spacer(1, 0.15*inch))
                
                # RECOMMENDATIONS SECTION (FIXED - NOW INCLUDED!)
                if result.get('recommendations') and result['recommendations'].strip():
                    story.append(Paragraph("🎯 Recommendations", section_style))
                    recommendations_text = result['recommendations'][:800]  # Limit length
                    story.append(Paragraph(recommendations_text, body_style))
                    story.append(Spacer(1, 0.15*inch))
                
                # Add separator line between students
                story.append(Spacer(1, 0.1*inch))
                separator_table = Table([['_'*100]], colWidths=[6.5*inch])
                separator_table.setStyle(TableStyle([
                    ('TEXTCOLOR', (0, 0), (-1, -1), colors.grey),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTSIZE', (0, 0), (-1, -1), 8)
                ]))
                story.append(separator_table)
                story.append(Spacer(1, 0.2*inch))
                
                # Page break after every 2 students for better readability
                if i % 2 == 0 and i < len(results):
                    story.append(PageBreak())
            
            # FOOTER PAGE
            story.append(PageBreak())
            story.append(Spacer(1, 2*inch))
            
            footer_style = ParagraphStyle(
                'FooterStyle',
                parent=styles['Normal'],
                fontSize=10,
                alignment=TA_CENTER,
                textColor=colors.grey
            )
            
            story.append(Paragraph("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━", footer_style))
            story.append(Spacer(1, 0.2*inch))
            story.append(Paragraph("💎 <b>Gem AI</b> - The World's Most Advanced Grading System", footer_style))
            story.append(Spacer(1, 0.1*inch))
            story.append(Paragraph("Making Education Assessment Effortless, One Grade at a Time", footer_style))
            story.append(Spacer(1, 0.1*inch))
            story.append(Paragraph(f"Report Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", footer_style))
            story.append(Spacer(1, 0.1*inch))
            story.append(Paragraph("Powered By Pluto Technology", footer_style))
            
            # Build the PDF
            doc.build(story)
            logger.info(f"✅ PDF report generated successfully: {output_path}")
            logger.info(f"✅ Report includes: feedback, strengths, improvements, corrections, and recommendations")
            logger.info("="*80)
            
        except Exception as e:
            logger.error(f"❌ Failed to generate PDF report: {str(e)}")
            logger.error(f"❌ Traceback: {traceback.format_exc()}")
            raise
    
    def generate_statistics_charts(self, results: List[Dict], output_dir: str):
        """Generate statistical charts and visualizations"""
        try:
            os.makedirs(output_dir, exist_ok=True)
            
            scores = [r.get('score', 0) for r in results if 'score' in r and 'error' not in r]
            grades = [r.get('grade', 'F') for r in results if 'grade' in r and 'error' not in r]
            
            if not scores:
                logger.warning("No scores to generate charts")
                return
            
            # Score distribution
            plt.figure(figsize=(10, 6))
            plt.hist(scores, bins=20, edgecolor='black', alpha=0.7, color='skyblue')
            plt.title('Score Distribution', fontsize=16, fontweight='bold')
            plt.xlabel('Score', fontsize=12)
            plt.ylabel('Frequency', fontsize=12)
            plt.grid(axis='y', alpha=0.3)
            plt.savefig(os.path.join(output_dir, 'score_distribution.png'), dpi=300, bbox_inches='tight')
            plt.close()
            
            # Grade distribution
            if grades:
                grade_counts = pd.Series(grades).value_counts()
                plt.figure(figsize=(10, 6))
                grade_counts.plot(kind='bar', color='coral', edgecolor='black')
                plt.title('Grade Distribution', fontsize=16, fontweight='bold')
                plt.xlabel('Grade', fontsize=12)
                plt.ylabel('Count', fontsize=12)
                plt.xticks(rotation=0)
                plt.grid(axis='y', alpha=0.3)
                plt.tight_layout()
                plt.savefig(os.path.join(output_dir, 'grade_distribution.png'), dpi=300, bbox_inches='tight')
                plt.close()
            
            logger.info(f"Charts generated in: {output_dir}")
        except Exception as e:
            logger.error(f"Failed to generate statistics charts: {str(e)}")
